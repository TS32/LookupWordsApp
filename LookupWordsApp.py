import os
import shutil
import traceback
import fstring
import pysnooper
import platform
import chardet
import requests
import urllib

from bs4 import BeautifulSoup

#from wordfreq import word_frequency
from tqdm import tqdm

from docx import Document
from docx.shared import Pt,Cm
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL


#@pysnooper.snoop(prefix='\n[class fileLoader]\t',thread_info=True)
class fileLoader():
    '''
    create a file list with absolute filename string list from a given folder with given extension.

    '''
    #@pysnooper.snoop(prefix='\n[class fileLoader __init__]\t',thread_info=True)
    def __init__(self,Path,Extension='.txt'):
        self.filepath=Path
        self.extension=Extension
        self.totalfiles=0
        self.relevantList=[]
        self.pathList=[]
        self.absoluteList=[]
        self.totalFileSize=0   

    #@pysnooper.snoop(prefix='\n[class fileLoader scanFiles]\t',thread_info=True)
    def scanFiles(self,maxfileCount=-1):
        '''
        scan the subfolders in the given path, looking for the files with given extension
        '''
        try:
            self.totalFileSize=0
            for parent, dirnames, filenames in os.walk(self.filepath, followlinks=False):
                num = 0   
                current_size_total=0         
                for filename in filenames:
                    if filename.endswith(self.extension):
                        self.relevantList.append(filename)
                        file_path = os.path.join(parent, filename)
                        self.absoluteList.append(file_path)
                        current_size_total += (os.stat(file_path).st_size)/1024
                        if parent not in self.pathList:
                            self.pathList.append(parent)
                        num += 1
                        if(maxfileCount!=-1 and num>=maxfileCount):
                            break
                if(num>0):        
                    print("\n\nfileLoader found %d files (totalFileSize=%d KB) with pattern [%s] in folder %s" % (num,current_size_total,self.extension,parent))
                    self.totalfiles += num  
                    # unit is KB
                    self.totalFileSize += current_size_total
                    if(maxfileCount!=-1 and self.totalfiles >=maxfileCount):
                        break
        except Exception as e:
            print("\nException:",e)
            print("\nCall Stack Trace:\n")
            traceback.print_exc()

        return self.totalfiles

#@pysnooper.snoop(prefix='\n[loadFileData]\t',thread_info=True)
def loadFileData(filename):
    data=None
    try:
        f=open(filename,"rb")
        raw_text=f.read()
        f.close()
        char_encoding=chardet.detect(raw_text)['encoding']
        print(f"\n\t  [Info]: file {filename} encoding type is {char_encoding}\n")
        data = raw_text.decode(char_encoding)
        if(len(data.strip())==0):
            return None
        else:
            return data
    except Exception as e:      
        print(f"\nException happened on loading data from {filename}:\n", e,"\n")  
        print("\nCall Stack Trace:\n")
        traceback.print_stack()
        return None

#@pysnooper.snoop(prefix='\n[generateWordList]\t',thread_info=True)
def generateWordList(filename):
    word_list=[]
    tempList=[]    

    fileContent = loadFileData(filename)    
    
    if fileContent is not None:
        tempList = fileContent.split("\n")
    
    for word in tempList:
        word = word.replace("\n", "").replace("\t", " ").replace("\r", "").replace("  "," ").strip()
        if(len(word)):
            word_list.append(word)
    
    print(f"\n\t  [Info]: {len(word_list)} words loaded! \n")
    
    if(len(word_list)):         
        return word_list
    else: 
        return None    
    
#@pysnooper.snoop(prefix='\n[lookupword]\t',thread_info=True)
def lookupword(word):
    # lookup single English word from dict-co.iciba by API call
    # API return value: key(单词),ps（音标）,pron（音频url）,pos（词性）,acceptation（释义）
    url = 'http://dict-co.iciba.com/api/dictionary.php?w={}&key=4EE27DDF668AD6501DCC2DC75B46851B'.format(urllib.parse.quote_plus(word))
    #print(url)
    try:
        resp = requests.get(url)
        resp.encoding = 'utf-8'
        if resp.status_code!=200:
            return None
        soup = BeautifulSoup(resp.text,'html.parser')
        key = soup.key.string
        ps = '[{}]'.format(soup.ps.string if soup.ps is not None else " ")
        pron = soup.pron.string if soup.pron is not None else " "
        pos_list = soup.select('pos')      
        acceptation_list = soup.select('acceptation')
        pos=[]
        acceptation=[]       
        example_sentences=soup.select('sent')
        if(len(pos_list)>0 and pos_list[0].string is not None):
            #word_freq= word_frequency(key,'en') 
            word_freq= 1          
            for i in range(0,len(pos_list)):
                pos.append(pos_list[i].string.replace('\n','').replace('\r',''))
                acceptation.append(acceptation_list[i].string.replace('\n','').replace('\r',''))
            #generate return value
            RET={"word":key,"word_freq":word_freq,"phonetic":ps,"pronouncation":pron,"pos":pos,"acceptation":acceptation,"sent":example_sentences}
            #print(RET)
        elif(len(acceptation_list)>0):  #this might be a phrase, get everything except the pos
            #word_freq= word_frequency(key,'en') 
            word_freq= 1 
            pos.append(" ")            
            for i in range(0,len(acceptation_list)):                
                acceptation.append(acceptation_list[i].string.replace('\n','').replace('\r',''))
            RET={"word":key,"word_freq":word_freq,"phonetic":ps,"pronouncation":pron,"pos":pos,"acceptation":acceptation,"sent":example_sentences} 
            #print(RET)  
        else:
            RET=None                
        return RET
    except:        
        traceback.print_exc()
        return None
        

#@pysnooper.snoop(prefix='\n[CreateTranslationDocument]\t',thread_info=True)        
def CreateWordListDocument(inputFile,outputFile=None,):
    basename="" 
    (folder_path, shortname) = os.path.split(inputFile)
    (basename, extension) = os.path.splitext(shortname)
    if(outputFile is None):
        filename_docx_string=basename+r'.docx'
        outputFile=os.path.join(folder_path,filename_docx_string) 

    #Process the text data, split the English and Chinese sentences/vocabulary
        
    word_list=generateWordList(inputFile)
    num=len(word_list)
    if(num==0): 
        print(f"\n[Error]: There is no vocabulary in the text file {inputFile}, abort! \n")
        return None
    
    title=basename+" "+ "word list"
    title=title.title()
    sent_title = basename+" "+ "sentences examples"
    sent_title = sent_title.title()
    sentences_list=[]
    Page_topMargin=Cm(1.27)
    Page_bottomMargin=Cm(1.27)
    Page_leftMargin=Cm(1.27)
    Page_rightMargin=Cm(1.27)

    #For A4 page orientation Landscape, if portrait just swap height and width

    A4_Page_Height = Cm(29.7)
    A4_Page_Width  =Cm(21) 

    #29.7-1.27*2 = 27.16
    #6.5*2+7*2 = 27
    Word_column_width=Cm(3.5)
    Phonetic_column_width=Cm(3.5)
    POS_column_width=Cm(2)
    # 21-1.27*2 -9 = 
    Acceptation_column_width=Cm(9.4)
    
    #Init the document
    document = Document()

    #init the font style
    style = document.styles['Normal']
    style.font.name = 'Tahoma'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')         
    style.font.size = Pt(12)

    #changing the page margins
    sections = document.sections
    
    for section in sections:
        section.orientation = WD_ORIENT.PORTRAIT  # set the page landscape or portrait         

        if(section.orientation == WD_ORIENT.LANDSCAPE):
            section.page_width = A4_Page_Height
            section.page_height = A4_Page_Width
        else:
            section.page_width = A4_Page_Width
            section.page_height = A4_Page_Height

        section.top_margin = Page_topMargin
        section.bottom_margin = Page_bottomMargin
        section.left_margin = Page_leftMargin
        section.right_margin = Page_rightMargin

    #Add The title
    p=document.add_heading(title, 0)    
        
    #Add the Table
    table = document.add_table(1, 4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    table.autofit=False
    table.allow_autofit = False

    HEADER_ROW=0                     
    hdr_cells=table.rows[HEADER_ROW].cells  #add table header
    hdr_cells[0].text="Word"       
    hdr_cells[0].width=Word_column_width             
    hdr_cells[1].text="Phonetic"    
    hdr_cells[1].width=Phonetic_column_width             
    hdr_cells[2].text="POS"  
    hdr_cells[2].width=POS_column_width                                                       
    hdr_cells[3].text="Acceptation"  
    hdr_cells[3].width=Acceptation_column_width
    for cell in hdr_cells:
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
    for word_i in tqdm(word_list,desc="Lookup words",total=len(word_list),unit="words"):
        lookup_result = lookupword(word_i)
        if lookup_result is not None:  
            sentences_list.append({"word":lookup_result['word'],"sent":lookup_result['sent']})          
            #first insert necessary rows
            pos_num = len(lookup_result['pos'])
            for ss in range(pos_num):
                table.add_row()
                current_row=len(table.rows)-1
                current_cells=table.rows[current_row].cells
                current_cells[0].text=lookup_result['word']    
                current_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                current_cells[0].width=Word_column_width
                current_cells[1].text=lookup_result['phonetic']
                current_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                current_cells[1].width=Phonetic_column_width
                current_cells[2].text=lookup_result['pos'][ss]
                current_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                current_cells[2].width=POS_column_width
                current_cells[3].text=lookup_result['acceptation'][ss]
                current_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                current_cells[3].width=Acceptation_column_width
                if(ss>=1):
                    for col in range(2): # first 2 columns merged seperately
                        up_cell=table.cell(current_row-1,col)
                        down_cell=table.cell(current_row,col)
                        down_cell.merge(up_cell)
                    table.cell(current_row,0).text=lookup_result['word']
                    table.cell(current_row,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    table.cell(current_row,0).width=Word_column_width
                    table.cell(current_row,1).text=lookup_result['phonetic']
                    table.cell(current_row,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    table.cell(current_row,1).width=Phonetic_column_width
    
    columns = table.columns
    columns[0].width=Word_column_width
    columns[1].width=Phonetic_column_width
    columns[2].width=POS_column_width
    columns[3].width=Acceptation_column_width

    #Add example sentences
    if len(sentences_list):

        #Add The title
        p=document.add_heading("\n\n"+sent_title, 0)   

        for index,result in enumerate(sentences_list):
            word = result['word']
            sentences = result['sent']
            if(sentences is not None and len(sentences)):
                p=document.add_heading(f"{index+1:>3}. "+word, 1)
                for s_i, sent in enumerate(sentences):
                    English = " ".join(sent.orig.string.splitlines())
                    Chinese = "".join(sent.trans.string.splitlines())
                 
                    Example = f"({(s_i+1):>2} ). " + English +"\n"
                    paragraph = document.add_paragraph(Example, style='List')  #添加英文句子
                    paragraph.add_run(Chinese)                
                    paragraph.paragraph_format.space_before = Pt(10)                     

    document.save(outputFile)
    print(f"\nVocabulary translation document {outputFile} generated successfully!\n")
    return num


if __name__ == "__main__":

    OSTYPE=platform.system()
    
    print(f"\n[Info]: -Run on {OSTYPE}-")
  
    absFilePath = os.path.abspath(__file__)
    RootFolder, scriptFilename = os.path.split(absFilePath)
    print(f"\n[Info]: Current script absolute filename {absFilePath}")
    print(f"\n[Info]: Current script absolute root path {RootFolder}")       

    #Generate Input folder path and Output folder
    
    InputFolder=os.path.join(RootFolder,"input")
    OutputFolder=os.path.join(RootFolder,"output")
    print(f"\n[Info]: Input folder path {InputFolder}")      
    print(f"\n[Info]: Output folder path {OutputFolder}")     

    txtFileFinder = fileLoader(InputFolder,Extension=".txt")   
    total_job = txtFileFinder.scanFiles()  #start to search for the input files

    print(f"\nStart {total_job} files jobs from input folder {InputFolder} \n")
    if(total_job==0):
        print("\nNothing to do, Exit !\n")
        exit(0)

    #Delete the output folder and create again (empty the output folder)

    if(os.path.exists(OutputFolder)):
        shutil.rmtree(OutputFolder, ignore_errors=True)
        print(f'\nDelete {OutputFolder} directory!\n')

    if(not os.path.isdir(OutputFolder)):
        os.makedirs(OutputFolder)
        print(f'\nCreate {OutputFolder} directory!\n')
   
    count=0
    for inputfile in txtFileFinder.absoluteList:
        (folder_path, shortname) = os.path.split(inputfile)
        (basename, extension) = os.path.splitext(shortname)
        filename_docx_string=basename+r'.docx'
        outputfile=os.path.join(OutputFolder,filename_docx_string) 
        if(CreateWordListDocument(inputfile,outputfile) is not None):
            count+=1
    
    print(f"\nDone! {count} files processed successfully! \n")